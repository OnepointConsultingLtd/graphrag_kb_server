import asyncio
from pathlib import Path

import pytest
from PIL import Image


def _make_docx(path: Path) -> Path:
    from docx import Document
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a sample paragraph for image extraction testing.")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# get_docx_image_path
# ---------------------------------------------------------------------------

class TestGetDocxImagePath:

    def test_returns_path_with_correct_extension(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import get_docx_image_path

        docx = _make_docx(tmp_path / "sample.docx")
        result = get_docx_image_path(docx, "png")
        assert result == tmp_path / "sample.png"

    def test_returns_none_for_missing_file(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import get_docx_image_path

        result = get_docx_image_path(tmp_path / "nonexistent.docx", "png")
        assert result is None

    def test_respects_image_format(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import get_docx_image_path

        docx = _make_docx(tmp_path / "sample.docx")
        result = get_docx_image_path(docx, "jpeg")
        assert result == tmp_path / "sample.jpeg"


# ---------------------------------------------------------------------------
# get_docx_image_by_path
# ---------------------------------------------------------------------------

class TestGetDocxImageByPath:

    def test_returns_none_when_image_does_not_exist(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import get_docx_image_by_path

        docx = _make_docx(tmp_path / "sample.docx")
        result = get_docx_image_by_path(docx, "png")
        assert result is None

    def test_returns_path_string_when_image_exists(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import get_docx_image_by_path

        docx = _make_docx(tmp_path / "sample.docx")
        image_path = tmp_path / "sample.png"
        Image.new("RGB", (10, 10), color="white").save(image_path)

        result = get_docx_image_by_path(docx, "png")
        assert result is not None
        assert result.endswith("/sample.png")
        assert result.startswith("/")


# ---------------------------------------------------------------------------
# _sync_extract_image_from_docx
# ---------------------------------------------------------------------------

class TestSyncExtractImageFromDocx:

    def test_creates_image_from_docx(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import _sync_extract_image_from_docx

        docx = _make_docx(tmp_path / "sample.docx")
        result = _sync_extract_image_from_docx(docx, "png")

        assert result is not None
        image_path = tmp_path / "sample.png"
        assert image_path.exists()
        img = Image.open(image_path)
        assert img.width > 0
        assert img.height > 0

    def test_skips_when_image_already_exists(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import _sync_extract_image_from_docx

        docx = _make_docx(tmp_path / "sample.docx")
        sentinel = Image.new("RGB", (1, 1), color="red")
        image_path = tmp_path / "sample.png"
        sentinel.save(image_path)
        mtime_before = image_path.stat().st_mtime

        _sync_extract_image_from_docx(docx, "png")

        assert image_path.stat().st_mtime == mtime_before

    def test_returns_none_and_writes_error_file_on_failure(self, tmp_path):
        from unittest.mock import patch
        from graphrag_kb_server.service.docs_image_extraction import _sync_extract_image_from_docx

        docx = _make_docx(tmp_path / "broken.docx")

        with patch("graphrag_kb_server.service.docs_image_extraction.Image") as mock_pil:
            mock_pil.open.side_effect = Exception("simulated PIL failure")
            result = _sync_extract_image_from_docx(docx, "png")

        assert result is None
        error_file = tmp_path / "broken.error.txt"
        assert error_file.exists()
        assert len(error_file.read_text()) > 0

    def test_skips_when_error_file_exists(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import _sync_extract_image_from_docx

        docx = _make_docx(tmp_path / "sample.docx")
        error_file = tmp_path / "sample.error.txt"
        error_file.write_text("previous error", encoding="utf-8")

        result = _sync_extract_image_from_docx(docx, "png")

        assert result is None
        assert not (tmp_path / "sample.png").exists()


# ---------------------------------------------------------------------------
# _extract_image_from_docx (async)
# ---------------------------------------------------------------------------

class TestExtractImageFromDocx:

    def test_async_wrapper_creates_image(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import _extract_image_from_docx

        docx = _make_docx(tmp_path / "sample.docx")
        result = asyncio.run(_extract_image_from_docx(docx, "png"))

        assert result is not None
        assert (tmp_path / "sample.png").exists()

    def test_returns_none_for_non_docx_extension(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import _extract_image_from_docx

        txt = tmp_path / "sample.txt"
        txt.write_text("hello")
        result = asyncio.run(_extract_image_from_docx(txt, "png"))
        assert result is None


# ---------------------------------------------------------------------------
# extract_images_from_docx (async, project-level)
# ---------------------------------------------------------------------------

class TestExtractImagesFromDocx:

    def _make_project(self, root: Path, names: list[str]) -> list[Path]:
        input_dir = root / "original_input"
        input_dir.mkdir(parents=True)
        paths = []
        for name in names:
            paths.append(_make_docx(input_dir / name))
        return paths

    def test_extracts_all_docx_in_project_folder(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import extract_images_from_docx

        self._make_project(tmp_path, ["doc1.docx", "doc2.docx"])
        results = asyncio.run(extract_images_from_docx(tmp_path, "png"))

        assert len(results) == 2
        for path in results:
            assert path is not None
            assert path.endswith(".png")

    def test_returns_empty_list_when_no_docx_files(self, tmp_path):
        from graphrag_kb_server.service.docs_image_extraction import extract_images_from_docx

        (tmp_path / "original_input").mkdir(parents=True)
        results = asyncio.run(extract_images_from_docx(tmp_path, "png"))
        assert results == []

    def test_skips_failed_extractions(self, tmp_path):
        from unittest.mock import patch
        from graphrag_kb_server.service.docs_image_extraction import extract_images_from_docx

        input_dir = tmp_path / "original_input"
        input_dir.mkdir(parents=True)
        _make_docx(input_dir / "good.docx")
        _make_docx(input_dir / "bad.docx")

        call_count = 0

        def fake_open(stream):
            nonlocal call_count
            call_count += 1
            if call_count == 1:
                raise Exception("simulated failure on first file")
            from PIL import Image as RealImage
            return RealImage.open(stream)

        with patch("graphrag_kb_server.service.docs_image_extraction.Image") as mock_pil:
            mock_pil.open.side_effect = fake_open
            results = asyncio.run(extract_images_from_docx(tmp_path, "png"))

        assert len(results) == 1
